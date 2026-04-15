import base64
import io
from dataclasses import dataclass
from typing import Literal

import fitz
from docx import Document
from openpyxl import load_workbook

from app.services.file_validation import FileCategory
from app.services.pdf_to_images import pdf_bytes_to_png_base64_list


@dataclass(frozen=True)
class LLMInput:
    """Prepared payload for the extraction LLM."""

    mode: Literal["vision", "text"]
    image_parts: list[tuple[str, str]] | None = None
    text: str | None = None


PDF_TEXT_MEANINGFUL_MIN_CHARS = 300


def _pdf_to_text(data: bytes, *, max_pages: int = 10) -> str:
    """
    Extract text from a PDF using PyMuPDF.
    Returns an empty string when text extraction yields nothing useful.
    """
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return ""

    try:
        n = min(doc.page_count or 0, max_pages)
        if n <= 0:
            return ""
        chunks: list[str] = []
        for i in range(n):
            page = doc.load_page(i)
            text = (page.get_text("text") or "").strip()
            if text:
                chunks.append(text)
        return "\n\n".join(chunks).strip()
    except Exception:
        return ""
    finally:
        doc.close()


def _docx_to_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            line = " | ".join(cells)
            if line.strip():
                lines.append(line)
    return "\n".join(lines)


def _xlsx_to_text(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        out: list[str] = []
        for sheet_name in wb.sheetnames:
            out.append(f"{sheet_name}:")
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                cells: list[str] = []
                for v in row:
                    if v is None:
                        cells.append("")
                    else:
                        s = str(v).strip().replace("\n", " ").replace(",", ";")
                        cells.append(s)
                if any(c for c in cells):
                    out.append(",".join(cells))
        return "\n".join(out)
    finally:
        wb.close()


def build_llm_input(data: bytes, category: FileCategory) -> LLMInput:
    if category == FileCategory.PDF:
        # Hybrid strategy:
        # 1) Try native text extraction (cheaper/faster if the PDF contains real text).
        # 2) Fallback to rasterization + vision for scanned PDFs.
        text = _pdf_to_text(data)
        if len(text) >= PDF_TEXT_MEANINGFUL_MIN_CHARS:
            return LLMInput(mode="text", image_parts=None, text=text)

        images_b64 = pdf_bytes_to_png_base64_list(data)
        parts = [(b64, "image/png") for b64 in images_b64]
        return LLMInput(mode="vision", image_parts=parts, text=None)

    if category == FileCategory.PNG:
        b64 = base64.b64encode(data).decode("ascii")
        return LLMInput(mode="vision", image_parts=[(b64, "image/png")], text=None)

    if category == FileCategory.JPEG:
        b64 = base64.b64encode(data).decode("ascii")
        return LLMInput(mode="vision", image_parts=[(b64, "image/jpeg")], text=None)

    if category == FileCategory.DOCX:
        text = _docx_to_text(data).strip()
        if not text:
            raise ValueError("Could not extract text from the Word document.")
        return LLMInput(mode="text", image_parts=None, text=text)

    if category == FileCategory.XLSX:
        text = _xlsx_to_text(data).strip()
        if not text:
            raise ValueError("Could not extract data from the Excel file.")
        return LLMInput(mode="text", image_parts=None, text=text)

    raise ValueError("Unsupported category for content preparation.")
